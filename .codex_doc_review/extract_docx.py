from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def iter_blocks(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def main() -> None:
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    doc = Document(source)
    blocks = []
    for index, block in enumerate(iter_blocks(doc), start=1):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            drawings = len(block._p.xpath(".//w:drawing")) + len(block._p.xpath(".//w:pict"))
            if text or drawings:
                blocks.append({
                    "index": index,
                    "type": "paragraph",
                    "style": block.style.name if block.style else None,
                    "text": text,
                    "drawings": drawings,
                })
        else:
            rows = []
            for row in block.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            blocks.append({"index": index, "type": "table", "rows": rows})

    media = []
    with zipfile.ZipFile(source) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/"):
                info = archive.getinfo(name)
                media.append({"name": name, "size": info.file_size})

    data = {
        "source": str(source),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "blocks": blocks,
        "media": media,
        "headers": [[p.text for p in section.header.paragraphs] for section in doc.sections],
        "footers": [[p.text for p in section.footer.paragraphs] for section in doc.sections],
    }
    output.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
